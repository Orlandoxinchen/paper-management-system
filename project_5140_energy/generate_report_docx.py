# -*- coding: utf-8 -*-
"""
5140 Midterm — Project Option 4: Generate Word report from energy_forecast.py output.
Run: python energy_forecast.py  then  python generate_report_docx.py
"""
import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "output")
FIG_DIR = os.path.join(OUT_DIR, "figures")
OUT_WORD = os.path.join(BASE_DIR, "Energy_Consumption_Forecasting_Report.docx")

def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Times New Roman"

    title = doc.add_heading("Project Option 4: Energy Consumption Forecasting (Time Series)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("ECON 5140 — Applied Econometrics Midterm Project")
    doc.add_paragraph()

    results = {}
    results_path = os.path.join(OUT_DIR, "results.json")
    if os.path.isfile(results_path):
        with open(results_path, "r", encoding="utf-8") as f:
            results = json.load(f)

    # 1. Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "This project forecasts household electric power consumption (Global_active_power, kW) using historical usage. "
        "Minute-level data is aggregated to daily level. The goal is to produce a 1-week (7-day) ahead forecast using "
        "time series models: ARIMA, ETS (Holt-Winters), and Prophet. We analyze trend and seasonality (daily, weekly), "
        "compare models with MAE, RMSE, and MAPE, and discuss implications for energy planning."
    )
    doc.add_paragraph()

    # 2. Data
    doc.add_heading("2. Data", level=1)
    doc.add_paragraph(
        f"The series has {results.get('n_obs', '—')} daily observations. "
        f"Training: {results.get('n_train', '—')} days; test (hold-out): {results.get('n_test', '—')} days. "
        f"Forecast horizon: {results.get('forecast_horizon_days', 7)} days (1 week) ahead. "
        f"Training mean consumption: {results.get('train_mean', '—')} kW; standard deviation: {results.get('train_std', '—')}."
    )
    doc.add_paragraph()

    # 3. Exploratory Analysis
    doc.add_heading("3. Exploratory Analysis: Trend and Seasonality", level=1)
    doc.add_paragraph(
        f"Augmented Dickey-Fuller test: statistic = {results.get('adf_statistic', '—')}, p-value = {results.get('adf_pvalue', '—')}. "
        + ("The series is stationary at 5%." if results.get("adf_pvalue", 1) < 0.05 else "The series is non-stationary; differencing is used in ARIMA.")
    )
    doc.add_paragraph("Decomposition (trend + seasonal + residual) and ACF/PACF:")
    if os.path.isfile(os.path.join(FIG_DIR, "fig1_series_decomposition.png")):
        doc.add_picture(os.path.join(FIG_DIR, "fig1_series_decomposition.png"), width=Inches(5.5))
    if os.path.isfile(os.path.join(FIG_DIR, "fig2_acf_pacf.png")):
        doc.add_picture(os.path.join(FIG_DIR, "fig2_acf_pacf.png"), width=Inches(5.5))
    doc.add_paragraph()

    # 4. Methods
    doc.add_heading("4. Methods", level=1)
    doc.add_paragraph(
        "Three models were fitted on the training set for a 7-day ahead forecast: "
        "(1) ARIMA(1,1,1); (2) ETS (Holt-Winters) with additive trend and seasonal (period=7 for weekly); "
        "(3) Prophet with yearly and weekly seasonality. Forecasts are compared to the held-out test set."
    )
    doc.add_paragraph()

    # 5. Results
    doc.add_heading("5. Results: Model Comparison (MAE, RMSE, MAPE)", level=1)
    doc.add_paragraph("Evaluation on the 7-day test set:")
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "RMSE"
    hdr[2].text = "MAE"
    hdr[3].text = "MAPE (%)"
    for i, (mod, rmse, mae, mape) in enumerate([
        ("ARIMA(1,1,1)", results.get("arima_rmse"), results.get("arima_mae"), results.get("arima_mape")),
        ("ETS (Holt-Winters)", results.get("ets_rmse"), results.get("ets_mae"), results.get("ets_mape")),
        ("Prophet", results.get("prophet_rmse"), results.get("prophet_mae"), results.get("prophet_mape")),
    ], start=1):
        row = table.rows[i].cells
        row[0].text = mod
        row[1].text = str(rmse) if rmse is not None else "—"
        row[2].text = str(mae) if mae is not None else "—"
        row[3].text = str(mape) if mape is not None else "—"
    doc.add_paragraph()
    doc.add_paragraph("1-week ahead forecast vs actual:")
    if os.path.isfile(os.path.join(FIG_DIR, "fig3_forecasts.png")):
        doc.add_picture(os.path.join(FIG_DIR, "fig3_forecasts.png"), width=Inches(5.5))
    doc.add_paragraph()

    # 6. Conclusion and Insights for Energy Planning
    doc.add_heading("6. Conclusion and Insights for Energy Planning", level=1)
    doc.add_paragraph(
        "We forecasted household energy consumption 7 days ahead using ARIMA, ETS, and Prophet. "
        "Weekly seasonality is relevant for daily data; ETS and Prophet capture it explicitly. "
        "Lower MAE/RMSE/MAPE indicate better accuracy for load forecasting. "
        "These forecasts support short-term energy planning, demand-side management, and grid balancing. "
        "Using the best-performing model (e.g., lowest MAPE) can improve operational decisions for utilities and households."
    )
    doc.add_paragraph()

    # Appendix: Code
    doc.add_heading("Appendix: Code", level=1)
    doc.add_paragraph("Analysis and figures were produced with energy_forecast.py (pandas, statsmodels, matplotlib, Prophet).")
    code_path = os.path.join(BASE_DIR, "energy_forecast.py")
    if os.path.isfile(code_path):
        with open(code_path, "r", encoding="utf-8") as f:
            code_text = f.read()
        doc.add_paragraph(code_text[:5500] + ("\n... (truncated)" if len(code_text) > 5500 else ""), style="No Spacing")

    doc.save(OUT_WORD)
    print("Report saved:", OUT_WORD)

if __name__ == "__main__":
    main()
